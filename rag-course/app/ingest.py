"""阶段 1 + 升级②：文档接入 —— 支持 Markdown / TXT / PDF / DOCX。

流水线：加载（load）→ 清洗（clean）→ 分块（chunk）→ 元数据（metadata）

用法（在 rag-course 目录下）：
    python app/ingest.py
    python app/ingest.py --chunk-size 300 --overlap 30
"""
from __future__ import annotations

import argparse
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path

DATA_DIR = Path(__file__).resolve().parent.parent / "data" / "sample"
SUPPORTED_SUFFIXES = {".md", ".markdown", ".txt", ".pdf", ".docx"}
DEFAULT_CHUNK_SIZE = 500
DEFAULT_OVERLAP = 50


# ---------- 数据模型 ----------


@dataclass
class RawDocument:
    """加载后的一页文本：PDF 按页返回（带页码），文本文件只有一页。"""

    text: str
    page: int | None = None


@dataclass
class Chunk:
    """一个可供检索的最小单元：文本 + 来源信息。"""

    text: str
    metadata: dict[str, object] = field(default_factory=dict)


# ---------- 1. 加载 ----------


def load_document(path: Path) -> list[RawDocument]:
    """按扩展名加载文档，返回一页或多页文本。"""
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        return [RawDocument(text=load_text(path))]
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".docx":
        return load_docx(path)
    raise ValueError(f"暂不支持 {suffix}，当前支持：{sorted(SUPPORTED_SUFFIXES)}")


def load_text(path: Path) -> str:
    """读取文本文件：优先 UTF-8，解码失败回退 GB18030（国内常见编码）。"""
    for encoding in ("utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码：{path.name}")


def load_pdf(path: Path) -> list[RawDocument]:
    """用 PyMuPDF 按页提取文字：一页一个 RawDocument，页码进元数据。"""
    import fitz  # pymupdf

    pages: list[RawDocument] = []
    with fitz.open(path) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text().strip()
            if text:
                pages.append(RawDocument(text=text, page=page_num))
    return pages


def load_docx(path: Path) -> list[RawDocument]:
    """用 python-docx 提取段落和表格；标题样式尽量转成 Markdown 标题。"""
    from docx import Document

    lines: list[str] = []
    doc = Document(path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style is not None else ""
        match = re.match(r"^(?:Heading|标题)\s*(\d)$", style)
        if match:
            lines.append("#" * int(match.group(1)) + " " + text)
        else:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return [RawDocument(text="\n".join(lines))]


# ---------- 2. 清洗 ----------

# 控制字符（\r 已单独处理，其余 0x00-0x1f 和 DEL 直接删掉）
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
# 三个以上连续空行压缩成两个
_MULTIPLE_BLANK_LINES = re.compile(r"\n{3,}")


def clean_text(text: str) -> str:
    """统一换行 → NFKC 归一化 → 去控制字符 → 去行尾空格 → 压缩空行。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = unicodedata.normalize("NFKC", text)
    text = _CONTROL_CHARS.sub("", text)
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = _MULTIPLE_BLANK_LINES.sub("\n\n", text)
    return text.strip()


# ---------- 3. 分块 ----------

# 递归切分的分隔符优先级：先按大段，再按行、句、逗号
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
# MULTILINE：让 ^ 匹配每一行的开头，才能识别文档任意位置的标题
_HEADING_RE = re.compile(r"^#{1,6}\s+(.*)$", re.MULTILINE)

# 受保护内容：代码块 / Markdown 表格——分块时整体保留，不拆散
_PROTECTED_PATTERNS = [
    re.compile(r"```[\s\S]*?```"),                          # 代码块
    re.compile(r"(?:^|\n)(?:[ \t]*\|[^\n]*\|[ \t]*\n)+"),   # Markdown 表格块
]


def _protect(text: str) -> tuple[str, list[str]]:
    """把受保护内容替换成占位符，返回 (占位文本, 原文列表)。"""
    protected: list[str] = []

    def _repl(match: re.Match) -> str:
        protected.append(match.group(0))
        return f"\x01P{len(protected) - 1}\x01"

    protected_text = text
    for pattern in _PROTECTED_PATTERNS:
        protected_text = pattern.sub(_repl, protected_text)
    return protected_text, protected


def _restore(text: str, protected: list[str]) -> str:
    """把占位符还原成原文。"""
    for i, original in enumerate(protected):
        text = text.replace(f"\x01P{i}\x01", original)
    return text


def chunk_document(
    text: str,
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP,
) -> list[Chunk]:
    """先保护表格/代码块（整体不拆散），再分块；有标题走 markdown，否则递归切分。"""
    protected_text, protected = _protect(text)
    if _HEADING_RE.search(protected_text):
        chunks = _markdown_split(protected_text, chunk_size, overlap)
    else:
        chunks = [
            Chunk(piece) for piece in _recursive_split(protected_text, chunk_size, overlap)
        ]
    for chunk in chunks:
        chunk.text = _restore(chunk.text, protected)
    return chunks


def _recursive_split(text: str, chunk_size: int, overlap: int) -> list[str]:
    """按分隔符优先级逐级切分，overlap 让相邻块共享边界上下文。"""
    pieces: list[str] = []
    rest = text.strip()
    while len(rest) > chunk_size:
        split_at = 0
        for sep in _SEPARATORS:
            if sep == "":
                split_at = chunk_size
                break
            pos = rest.rfind(sep, 0, chunk_size)
            if pos > 0:
                split_at = pos + len(sep)
                break
        if split_at <= 0:  # 找不到任何分隔符，只能硬切
            split_at = chunk_size
        pieces.append(rest[:split_at].strip())
        # 重叠区域最多到 split_at-1，保证剩余部分每次都在缩短（防止死循环）
        cut = max(0, split_at - overlap)
        if cut <= 0:
            cut = split_at
        rest = rest[cut:]
    if rest.strip():
        pieces.append(rest.strip())
    return pieces


def _extract_sections(text: str) -> list[tuple[str, str]]:
    """按标题层级切出 (章节路径, 正文)，正文不含标题行。"""
    sections: list[tuple[str, str]] = []
    current_lines: list[str] = []
    current_path = ""
    heading_stack: list[tuple[int, str]] = []

    def flush() -> None:
        nonlocal current_lines
        body = "\n".join(current_lines).strip()
        if body:
            sections.append((current_path, body))
        current_lines = []

    for line in text.split("\n"):
        match = _HEADING_RE.match(line)
        if match:
            flush()
            level = len(line) - len(line.lstrip("#"))
            title = match.group(1).strip()
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, title))
            current_path = " / ".join(t for _, t in heading_stack)
        else:
            current_lines.append(line)
    flush()
    return sections


def _markdown_split(text: str, chunk_size: int, overlap: int) -> list[Chunk]:
    """按标题分块：同父标题的小节可合并到接近 chunk_size，超长小节单独递归切分。

    metadata.section = 第一个小节路径（显示用）；metadata.sections = 覆盖的所有小节路径
    （标题上下文：合并块引用时不会指错小节）。
    """
    chunks: list[Chunk] = []
    buffer: list[str] = []
    buffer_path = ""
    buffer_sections: list[str] = []

    def flush_buffer() -> None:
        nonlocal buffer, buffer_path, buffer_sections
        if buffer:
            chunks.append(
                Chunk(
                    "\n\n".join(buffer),
                    {"section": buffer_path, "sections": buffer_sections},
                )
            )
        buffer, buffer_path, buffer_sections = [], "", []

    for path, body in _extract_sections(text):
        if len(body) > chunk_size:
            flush_buffer()
            for piece in _recursive_split(body, chunk_size, overlap):
                chunks.append(Chunk(piece, {"section": path, "sections": [path]}))
            continue
        # 父标题不同的小节不允许合并，防止跨章串味
        parent = path.rpartition(" / ")[0]
        buffer_parent = buffer_path.rpartition(" / ")[0] if buffer_path else ""
        if buffer and (parent != buffer_parent or len("\n\n".join(buffer)) + len(body) + 2 > chunk_size):
            flush_buffer()
        buffer.append(body)
        buffer_path = buffer_path or path
        if path not in buffer_sections:
            buffer_sections.append(path)
    flush_buffer()
    return chunks


# ---------- 4. 流水线（加载 + 清洗 + 分块 + 元数据） ----------


def process_document(
    path: Path,
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_OVERLAP,
) -> list[Chunk]:
    chunks: list[Chunk] = []
    index = 0  # 跨页全局递增，保证每份文档的块 id 唯一
    for raw in load_document(path):
        cleaned = clean_text(raw.text)
        if not cleaned:
            continue
        for chunk in chunk_document(cleaned, chunk_size=chunk_size, overlap=overlap):
            chunk.metadata.update(
                {
                    "file_name": path.name,
                    "file_path": str(path),
                    "page": raw.page,
                    "chunk_index": index,
                }
            )
            chunks.append(chunk)
            index += 1
    return chunks


def main() -> None:
    parser = argparse.ArgumentParser(description="文档接入流水线：加载 → 清洗 → 分块 → 元数据")
    parser.add_argument("--data-dir", type=Path, default=DATA_DIR)
    parser.add_argument("--chunk-size", type=int, default=DEFAULT_CHUNK_SIZE)
    parser.add_argument("--overlap", type=int, default=DEFAULT_OVERLAP)
    args = parser.parse_args()

    files = sorted(
        p for p in args.data_dir.iterdir() if p.suffix.lower() in SUPPORTED_SUFFIXES
    )
    if not files:
        print(f"目录中没有支持的文档（{sorted(SUPPORTED_SUFFIXES)}）：{args.data_dir}")
        sys.exit(1)

    total = 0
    for path in files:
        chunks = process_document(path, chunk_size=args.chunk_size, overlap=args.overlap)
        total += len(chunks)
        chars = sum(len(c.text) for c in chunks)
        avg = chars // len(chunks) if chunks else 0
        print(f"[{path.name}] 块数={len(chunks)} 总字符={chars} 平均={avg}")
        if chunks:
            print(f"  首块预览: {chunks[0].text[:70].replace(chr(10), ' ')}…")
            print(f"  元数据: {chunks[0].metadata}")
    print(f"\n共 {len(files)} 个文档、{total} 个块（chunk_size={args.chunk_size}, overlap={args.overlap}）")


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass
    main()
